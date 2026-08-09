from __future__ import annotations

import csv
import json
import math
import os
import shutil
import textwrap
import zipfile
from pathlib import Path

import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import font_manager
from graphviz import Digraph
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCRIPT = Path(__file__).resolve()
BASE = SCRIPT.parents[2]
OUT = SCRIPT.parent
ASSETS = OUT / 'assets'
ANNEX = OUT / 'annexes'
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)
ANNEX.mkdir(parents=True, exist_ok=True)

# ---------- Data ----------
tr = pd.read_csv(BASE / '01_branche_matiere/base_transitions/transitions_matiere.csv', sep=';', encoding='utf-8-sig')
ga = pd.read_csv(BASE / '00_socle/genealogie/arbre_genealogique.csv', sep=';', encoding='utf-8-sig')
gm = pd.read_csv(BASE / '01_branche_matiere/genealogie/genealogie_matiere.csv', sep=';', encoding='utf-8-sig')
with open(BASE / '00_socle/genealogie/cloture_arbre.json', encoding='utf-8') as f:
    close_ga = json.load(f)
with open(BASE / '01_branche_matiere/genealogie/cloture_genealogie.json', encoding='utf-8') as f:
    close_gm = json.load(f)
with open(BASE / '01_branche_matiere/hypergraphe_transformations/validation_hypergraphe.json', encoding='utf-8') as f:
    hyper_valid = json.load(f)
with open(BASE / '01_branche_matiere/hypergraphe_transformations/test_hierarchie_resultats.json', encoding='utf-8') as f:
    hierarchy = json.load(f)
with open(BASE / '01_branche_matiere/hypergraphe_transformations/inventaire_accessible_resultats.json', encoding='utf-8') as f:
    inventory = json.load(f)
with open(BASE / '01_branche_matiere/memoire_materielle_reelle/derive/execution/CAMPAGNE.json', encoding='utf-8') as f:
    material_memory = json.load(f)

# Copies utiles dans les annexes du document scientifique.
for rel in [
    '00_socle/genealogie/arbre_genealogique.csv',
    '00_socle/genealogie/cloture_arbre.json',
    '00_socle/genealogie/correspondance_GM_GA.csv',
    '00_socle/genealogie/REFERENCES_TRANSITIONS.csv',
    '01_branche_matiere/genealogie/PROTOCOLE_INFORMATION_GENEALOGIQUE.md',
    '01_branche_matiere/genealogie/genealogie_matiere.csv',
    '01_branche_matiere/genealogie/cloture_genealogie.json',
    '01_branche_matiere/hypergraphe_transformations/inventaire_accessible_resultats.json',
    '01_branche_matiere/hypergraphe_transformations/test_hierarchie_resultats.json',
    '02_branche_systeme_solaire/couche_astronomique/STATUT_SCIENTIFIQUE.md',
    '02_branche_systeme_solaire/couche_memoire_historique/REPORT.md',
    '03_branche_vivant/programme_prebiotique/PROGRAMME_PREBIOTIQUE.md',
]:
    src_path = BASE / rel
    if src_path.exists():
        shutil.copy2(src_path, ANNEX / src_path.name)

# ---------- Helpers for figures ----------
PURPLE = '#6C4CF1'
NAVY = '#243B67'
TEAL = '#138A72'
ORANGE = '#D97706'
RED = '#C0392B'
GRAY = '#667085'
LIGHT = '#F7F8FC'


def gnode(dot, node_id, label, color, dashed=False, fill='white', fontsize='9'):
    dot.node(node_id, label=label, shape='box', style=('rounded,dashed,filled' if dashed else 'rounded,filled'),
             color=color, fillcolor=fill, fontname='Arial', fontsize=fontsize, penwidth='1.3', margin='0.08,0.05')


def render_graph(dot: Digraph, name: str, fmt='png') -> Path:
    path = ASSETS / name
    cible = ASSETS / f'{path.stem}.{fmt}'
    if shutil.which('dot') is None:
        if cible.exists():
            print(f'Graphviz absent : figure versionnée réutilisée — {cible.name}')
            return cible
        raise RuntimeError(f'Graphviz absent et figure introuvable : {cible}')
    dot.render(filename=path.stem, directory=str(ASSETS), format=fmt, cleanup=True)
    return cible


def render_chain_figure() -> Path:
    """Rend la chaîne canonique sans dépendre de l'exécutable Graphviz."""
    labels = [
        'Histoire\nHₜ', 'Mémoire\nmₜ^ℓ', 'État et contraintes\nSₜ^ℓ, Θeff,t^ℓ',
        'P^adm', 'P^att', 'P^kin', 'P^pers', 'Réalisation\nRéalₜ₊₁',
    ]
    colors = [NAVY, NAVY, PURPLE, PURPLE, ORANGE, ORANGE, TEAL, RED]
    fig, ax = plt.subplots(figsize=(15.5, 2.25))
    ax.set_xlim(-0.5, len(labels) - 0.5)
    ax.set_ylim(-0.6, 0.6)
    ax.axis('off')
    for i, (label, color) in enumerate(zip(labels, colors)):
        ax.text(i, 0, label, ha='center', va='center', fontsize=9,
                bbox=dict(boxstyle='round,pad=0.42', facecolor='white',
                          edgecolor=color, linewidth=1.4))
        if i:
            ax.annotate('', xy=(i - 0.43, 0), xytext=(i - 0.57, 0),
                        arrowprops=dict(arrowstyle='->', color=GRAY, lw=1.2))
    cible = ASSETS / '03_chaine_oric.png'
    fig.savefig(cible, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return cible

# Figure 1: Programme architecture
D = Digraph('program', graph_attr={'rankdir':'TB','bgcolor':'white','pad':'0.25','nodesep':'0.35','ranksep':'0.55'})
D.attr('node', fontname='Arial')
gnode(D,'socle','SOCLE COMMUN ORI-C\narchitecture, mémoire, persistance, preuve, possibles',PURPLE,fill='#F0EDFF',fontsize='13')
for n,l,c,fill in [
    ('m','BRANCHE 1 - MATIÈRE\nrégimes 1 à 4',PURPLE,'#F4F1FF'),
    ('p','BRANCHE 2 - PLANÈTES ET SYSTÈME SOLAIRE\nrégimes 5 et 6',ORANGE,'#FFF5E8'),
    ('v','BRANCHE 3 - VIVANT\nrégimes 7 et 8',TEAL,'#EAF8F4')]:
    gnode(D,n,l,c,fill=fill,fontsize='11')
D.edge('socle','m',label='vocabulaire et critères',color=GRAY,fontname='Arial',fontsize='8')
D.edge('socle','p',label='vocabulaire et critères',color=GRAY,fontname='Arial',fontsize='8')
D.edge('socle','v',label='vocabulaire et critères',color=GRAY,fontname='Arial',fontsize='8')
D.edge('m','p',label='héritage matériel',color=PURPLE,fontname='Arial',fontsize='9')
D.edge('p','v',label='conditions planétaires',color=ORANGE,fontname='Arial',fontsize='9')
D.edge('v','p',label='rétroactions biogéochimiques',color=TEAL,fontname='Arial',fontsize='9',constraint='false')
program_fig=render_graph(D,'01_architecture_programme')

# Figure 2: Base layers
D = Digraph('layers', graph_attr={'rankdir':'TB','bgcolor':'white','pad':'0.25','nodesep':'0.22','ranksep':'0.28'})
layers=[
('L1','1. ARCHITECTURE PRÉSENTE','A(t) = [n, G, I, E, Π, H]','Composition, configuration, interactions, environnement, persistance, histoire',PURPLE,'#F2EFFF'),
('L2','2. DYNAMIQUE','Sₜ^ℓ, mₜ^ℓ, Aₜ^ℓ','L’état, les mémoires et l’opérateur d’évolution changent séparément à l’échelle déclarée',NAVY,'#EEF3FB'),
('L3','3. TRANSITION','S = (ΔV, ΔC, ΔΠ, ΔH, ΔR, ΔF)','Variables collectives, connectivité, persistance, héritage, robustesse, fermetures',ORANGE,'#FFF6EA'),
('L4','4. GÉNÉALOGIE','parents + conditions → mécanisme → produit','Les produits antérieurs deviennent les ressources matérielles des étapes suivantes',TEAL,'#EAF8F4'),
('L5','5. PREUVE','mécanisme, nature, histoire, rôle causal','Les niveaux de preuve et les tests restent séparés par domaine',RED,'#FFF0EF')]
for i,(nid,title,formula,desc,color,fill) in enumerate(layers):
    gnode(D,nid,f'{title}\n{formula}\n{desc}',color,fill=fill,fontsize='10')
    if i: D.edge(layers[i-1][0],nid,color=GRAY)
base_layers_fig=render_graph(D,'02_couches_architecture_scientifique')

# Figure 3: chaîne ORI-C, rendue sans dépendance Graphviz afin que le dossier
# reste reconstructible sur Windows à partir d'un clone standard.
chain_fig=render_chain_figure()

# Figure 4: Four representations
D=Digraph('repr',graph_attr={'rankdir':'LR','bgcolor':'white','pad':'0.25','nodesep':'0.45','ranksep':'0.3'})
reprs=[
('TR','CARTE DES TRANSITIONS\n40 transitions\n47 relations typées\nfonction documentaire',NAVY,'#EEF3FB'),
('HG','HYPERGRAPHE MÉCANISTIQUE\n53 nœuds\n53 hyperarêtes\nprocessus multi-entrées',PURPLE,'#F2EFFF'),
('GM','GÉNÉALOGIE DÉTAILLÉE\n22 transitions matière\n38 relations internes\ndétail de la matière',ORANGE,'#FFF6EA'),
('GA','ARBRE GÉNÉALOGIQUE GLOBAL\n39 transitions\n3 branches\n77 relations internes',TEAL,'#EAF8F4')]
for n,l,c,f in reprs: gnode(D,n,l,c,fill=f,fontsize='11')
D.edge('TR','HG',label='reclassement mécanistique',fontname='Arial',fontsize='8',color=GRAY)
D.edge('HG','GM',label='filiation explicite',fontname='Arial',fontsize='8',color=GRAY)
D.edge('GM','GA',label='raccordement des branches',fontname='Arial',fontsize='8',color=GRAY)
repr_fig=render_graph(D,'04_quatre_representations')

# Tree figures
certainty_dash={'plausible':True,'fortement inféré':True,'établi':False}
branch_color={'1 matière':PURPLE,'2 planétaire':ORANGE,'3 vivant':TEAL}
branch_fill={'1 matière':'#F6F3FF','2 planétaire':'#FFF7EC','3 vivant':'#ECF9F5'}

def product_short(s, width=22):
    s=s.replace(' | ', ', ')
    return textwrap.shorten(s,width=width,placeholder='…')

# overview, use clustered ranks
D=Digraph('tree',graph_attr={'rankdir':'LR','bgcolor':'white','pad':'0.15','nodesep':'0.16','ranksep':'0.32','splines':'spline'})
D.attr('edge',arrowsize='0.55',penwidth='0.8',color='#8991A3')
for branch in ['1 matière','2 planétaire','3 vivant']:
    with D.subgraph(name='cluster_'+branch.split()[0]) as c:
        c.attr(label=branch.upper(),color=branch_color[branch],fontcolor=branch_color[branch],fontname='Arial',fontsize='12',style='rounded')
        subset=ga[ga['branche']==branch]
        for _,r in subset.iterrows():
            label=f"{r['id']}\n{product_short(r['produit'],24)}"
            gnode(c,r['id'],label,branch_color[branch],dashed=certainty_dash.get(r['degre_de_certitude'],False),fill=branch_fill[branch],fontsize='8')
# internal direct sequential edges
for _,r in ga.iterrows():
    nxt=str(r['transition_suivante'])
    if nxt and nxt!='nan' and nxt in set(ga['id']):
        D.edge(r['id'],nxt)
# cross material edges
produced_by={}
for _,r in ga.iterrows():
    for p in [x.strip() for x in str(r['produit']).split('|') if x.strip()]: produced_by[p]=r['id']
for _,r in ga.iterrows():
    for p in [x.strip() for x in str(r['parents_materiels']).split('|') if x.strip()]:
        src=produced_by.get(p)
        if src and src!=r['id'] and ga.loc[ga['id']==src,'branche'].iloc[0]!=r['branche']:
            D.edge(src,r['id'],color=RED,penwidth='1.0',constraint='false')
tree_fig=render_graph(D,'05_arbre_genealogique_vue_ensemble')

# Branch diagrams
branch_figs={}
for branch, fname, title in [('1 matière','06_branche_matiere','Branche 1 - architectures matérielles'),('2 planétaire','07_branche_planetaire','Branche 2 - architectures planétaires'),('3 vivant','08_branche_vivant','Branche 3 - architectures du vivant')]:
    D=Digraph(branch,graph_attr={'rankdir':'LR','bgcolor':'white','pad':'0.2','nodesep':'0.2','ranksep':'0.25','label':title,'labelloc':'t','fontname':'Arial','fontsize':'15','fontcolor':branch_color[branch]})
    D.attr('edge',arrowsize='0.65',penwidth='0.9',color='#8991A3')
    subset=ga[ga['branche']==branch]
    ids=set(subset['id'])
    for _,r in subset.iterrows():
        label=f"{r['id']}\n{product_short(r['produit'],30)}\n{r['mecanisme_categorie'].replace('_',' ')}"
        col=branch_color[branch]
        fill=branch_fill[branch]
        gnode(D,r['id'],label,col,dashed=certainty_dash.get(r['degre_de_certitude'],False),fill=fill,fontsize='8')
    for _,r in subset.iterrows():
        nxt=str(r['transition_suivante'])
        if nxt in ids: D.edge(r['id'],nxt)
    branch_figs[branch]=render_graph(D,fname)


# Remplacement des schémas horizontaux par des versions paginables et lisibles.
FONT_REG = font_manager.findfont(font_manager.FontProperties(family='DejaVu Sans'))
FONT_BOLD = font_manager.findfont(
    font_manager.FontProperties(family='DejaVu Sans', weight='bold')
)


def _rgb(hex_color):
    h=hex_color.lstrip('#')
    return tuple(int(h[i:i+2],16) for i in (0,2,4))


def _wrap(draw, text, font, max_width):
    words=str(text).replace(' | ', ', ').split()
    lines=[]; current=''
    for word in words:
        candidate=(current+' '+word).strip()
        if draw.textbbox((0,0),candidate,font=font)[2] <= max_width:
            current=candidate
        else:
            if current: lines.append(current)
            current=word
    if current: lines.append(current)
    return lines


def _arrow(draw, p1, p2, color=(110,120,140), width=5):
    import math as _math
    draw.line([p1,p2],fill=color,width=width)
    ang=_math.atan2(p2[1]-p1[1],p2[0]-p1[0])
    size=15
    for delta in (2.55,-2.55):
        q=(p2[0]+size*_math.cos(ang+delta),p2[1]+size*_math.sin(ang+delta))
        draw.line([p2,q],fill=color,width=width)


def draw_branch_flow(branch, path, title, cols):
    subset=ga[ga['branche']==branch].reset_index(drop=True)
    n=len(subset); rows=math.ceil(n/cols)
    box_w, box_h, gap_x, gap_y = 330, 150, 55, 55
    left, top, bottom = 70, 155, 70
    W=left*2 + cols*box_w + (cols-1)*gap_x
    H=top + rows*box_h + (rows-1)*gap_y + bottom
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    title_size=38
    title_font=ImageFont.truetype(FONT_BOLD,title_size)
    while d.textbbox((0,0),title,font=title_font)[2] > W-80 and title_size > 22:
        title_size -= 2
        title_font=ImageFont.truetype(FONT_BOLD,title_size)
    id_font=ImageFont.truetype(FONT_BOLD,23)
    txt_font=ImageFont.truetype(FONT_REG,18)
    small_font=ImageFont.truetype(FONT_BOLD,14)
    d.text((W//2,35),title,font=title_font,fill=_rgb(branch_color[branch]),anchor='ma')
    positions=[]
    for i in range(n):
        r=i//cols; c=i%cols
        if r%2==1: c=cols-1-c
        x=left+c*(box_w+gap_x); y=top+r*(box_h+gap_y)
        positions.append((x,y))
    # arrows first
    for i in range(n-1):
        x1,y1=positions[i]; x2,y2=positions[i+1]
        if y1==y2:
            if x2>x1: p1=(x1+box_w,y1+box_h//2); p2=(x2,y2+box_h//2)
            else: p1=(x1,y1+box_h//2); p2=(x2+box_w,y2+box_h//2)
        else:
            p1=(x1+box_w//2,y1+box_h); p2=(x2+box_w//2,y2)
        _arrow(d,p1,p2)
    # boxes
    for i,(_,row) in enumerate(subset.iterrows()):
        x,y=positions[i]
        certainty=str(row['degre_de_certitude'])
        outline=_rgb(branch_color[branch])
        fill=_rgb(branch_fill[branch])
        d.rounded_rectangle((x,y,x+box_w,y+box_h),radius=20,fill=fill,outline=outline,width=5)
        d.text((x+16,y+12),str(row['id']),font=id_font,fill=outline)
        tag='PLAUSIBLE' if certainty=='plausible' else ('INFÉRÉ' if certainty=='fortement inféré' else 'ÉTABLI')
        d.text((x+box_w-14,y+16),tag,font=small_font,fill=outline,anchor='ra')
        prod_lines=_wrap(d,str(row['produit']),txt_font,box_w-30)[:2]
        mech_lines=_wrap(d,str(row['mecanisme_categorie']).replace('_',' '),txt_font,box_w-30)[:2]
        yy=y+54
        for line in prod_lines:
            d.text((x+16,yy),line,font=txt_font,fill=(25,30,40)); yy+=23
        yy=max(yy,y+108)
        for line in mech_lines:
            d.text((x+16,yy),line,font=txt_font,fill=(80,85,100)); yy+=21
    im.save(path,dpi=(180,180))
    return path


branch_figs['1 matière']=draw_branch_flow('1 matière',ASSETS/'06_branche_matiere.png','BRANCHE 1 - ARCHITECTURES MATÉRIELLES',3)
branch_figs['2 planétaire']=draw_branch_flow('2 planétaire',ASSETS/'07_branche_planetaire.png','BRANCHE 2 - ARCHITECTURES PLANÉTAIRES',2)
branch_figs['3 vivant']=draw_branch_flow('3 vivant',ASSETS/'08_branche_vivant.png','BRANCHE 3 - ARCHITECTURES DU VIVANT',2)


def draw_tree_overview(path):
    W,H=1500,1900
    im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    title_font=ImageFont.truetype(FONT_BOLD,44)
    head_font=ImageFont.truetype(FONT_BOLD,27)
    id_font=ImageFont.truetype(FONT_BOLD,16)
    text_font=ImageFont.truetype(FONT_REG,14)
    d.text((W//2,35),'ARBRE GÉNÉALOGIQUE GLOBAL',font=title_font,fill=_rgb(NAVY),anchor='ma')
    panels=[('1 matière',115,560,6),('2 planétaire',670,420,5),('3 vivant',1090,420,5)]
    pos={}
    for branch,y,panel_h,cols in panels:
        color=_rgb(branch_color[branch]); fill=_rgb(branch_fill[branch])
        d.rounded_rectangle((45,y,1455,y+panel_h),radius=24,fill=fill,outline=color,width=5)
        label=f"{branch.replace('1 ', '').replace('2 ', '').replace('3 ', '').upper()} - {len(ga[ga['branche']==branch])} TRANSITIONS"
        d.text((75,y+20),label,font=head_font,fill=color)
        subset=ga[ga['branche']==branch].reset_index(drop=True)
        rows=math.ceil(len(subset)/cols)
        bw=210 if cols==6 else 248; bh=105; gx=18; gy=22
        total_w=cols*bw+(cols-1)*gx
        x0=(W-total_w)//2; yy0=y+75
        seq_positions=[]
        for i,(_,row) in enumerate(subset.iterrows()):
            rr=i//cols; cc=i%cols
            if rr%2==1: cc=cols-1-cc
            x=x0+cc*(bw+gx); yy=yy0+rr*(bh+gy)
            seq_positions.append((x,yy))
            pos[row['id']]=(x+bw//2,yy+bh//2)
        for i in range(len(seq_positions)-1):
            x1,y1=seq_positions[i]; x2,y2=seq_positions[i+1]
            if y1==y2:
                p1=(x1+bw,y1+bh//2) if x2>x1 else (x1,y1+bh//2)
                p2=(x2,y2+bh//2) if x2>x1 else (x2+bw,y2+bh//2)
            else:
                p1=(x1+bw//2,y1+bh); p2=(x2+bw//2,y2)
            _arrow(d,p1,p2,color=(145,150,165),width=3)
        for i,(_,row) in enumerate(subset.iterrows()):
            x,yy=seq_positions[i]
            ol=_rgb(branch_color[branch])
            fi=_rgb('#FFFFFF')
            d.rounded_rectangle((x,yy,x+bw,yy+bh),radius=14,fill=fi,outline=ol,width=3)
            d.text((x+10,yy+8),row['id'],font=id_font,fill=ol)
            lines=_wrap(d,str(row['produit']),text_font,bw-18)[:3]
            ty=yy+36
            for line in lines:
                d.text((x+9,ty),line,font=text_font,fill=(35,40,50)); ty+=18
    # interbranch summary at bottom
    d.rounded_rectangle((75,1550,1425,1840),radius=22,fill=(250,250,252),outline=_rgb(RED),width=4)
    d.text((100,1575),'8 COUPLES DE TRANSITIONS INTERBRANCHES',font=head_font,fill=_rgb(RED))
    items=[
        'GA-019 → GA-020 : agrégats poreux vers galets',
        'GA-012 → GA-023 : aluminium 26 vers chauffage précoce',
        'GA-015 → GA-027 : glaces vers atmosphère et hydrosphère',
        'GA-017 → GA-029 : condensats réfractaires vers minéraux planétaires',
        'GA-016 → GA-030 : organiques complexes vers monomères',
        'GA-016 → GA-033 : organiques complexes vers vésicules',
        'GA-029 → GA-030 : minéraux et gradients vers monomères',
        'GA-029 → GA-037 : gradients redox vers métabolismes codés']
    for j,item in enumerate(items):
        col=0 if j<4 else 1; rr=j if j<4 else j-4
        x=105+col*650; yy=1635+rr*45
        d.ellipse((x,yy+5,x+13,yy+18),fill=_rgb(RED))
        d.text((x+25,yy),item,font=text_font,fill=(35,40,50))
    im.save(path,dpi=(180,180))
    return path


tree_fig=draw_tree_overview(ASSETS/'05_arbre_genealogique_vue_ensemble.png')

# Interbranch diagram
D=Digraph('inter',graph_attr={'rankdir':'LR','bgcolor':'white','pad':'0.3','nodesep':'0.4','ranksep':'0.45'})
D.attr('edge',fontname='Arial',fontsize='8',arrowsize='0.7')
for node,label,color,fill in [('M','BRANCHE MATIÈRE',PURPLE,'#F2EFFF'),('P','BRANCHE PLANÉTAIRE',ORANGE,'#FFF6EA'),('V','BRANCHE VIVANT',TEAL,'#EAF8F4')]:
    gnode(D,node,label,color,fill=fill,fontsize='13')
transfers=[
('M','P','agrégats poreux → galets'),
('M','P','²⁶Al → chauffage précoce'),
('M','P','glaces → atmosphère et hydrosphère'),
('M','P','condensats → minéraux planétaires'),
('M','V','molécules organiques → monomères'),
('M','V','molécules organiques → vésicules'),
('P','V','minéraux et gradients → monomères'),
('P','V','gradients redox → métabolismes')]
for i,(a,b,l) in enumerate(transfers):
    mid=f'x{i}'
    gnode(D,mid,l,GRAY,fill='white',fontsize='8')
    D.edge(a,mid,color=branch_color['1 matière'] if a=='M' else branch_color['2 planétaire'])
    D.edge(mid,b,color=RED)
inter_fig=render_graph(D,'09_transferts_interbranches')

# Corrected Al26 lineage
D=Digraph('al26',graph_attr={'rankdir':'LR','bgcolor':'white','pad':'0.3','nodesep':'0.3','ranksep':'0.3'})
D.attr('edge',arrowsize='0.7',color=GRAY)
steps=[
('a','Mg et Al dans les étoiles massives',PURPLE,'#F2EFFF'),
('b','Combustion de H\ncoquille C\nexplosif Ne/C',ORANGE,'#FFF6EA'),
('c','²⁶Al',RED,'#FFF0EF'),
('d','Éjection par vents\nou explosion',NAVY,'#EEF3FB'),
('e','Nuage et disque\nprotosolaires',NAVY,'#EEF3FB'),
('f','Incorporation\naux solides',PURPLE,'#F2EFFF'),
('g','Chauffage précoce\ndes planétésimaux',ORANGE,'#FFF6EA')]
for n,l,c,f in steps: gnode(D,n,l,c,fill=f,fontsize='10')
for (a,*_),(b,*__) in zip(steps,steps[1:]): D.edge(a,b)
al26_fig=render_graph(D,'10_correction_aluminium_26')

# Results status chart (matplotlib)
fig, ax = plt.subplots(figsize=(10,5.2))
labels=['Socle\nchémostat','Carte\nrelationnelle','Capacités\nmatière','Inventaire\naccessible','Astronomie\nN-corps','Mémoire\nM2','Exoplanète\npersistance','Prébiotique']
values=[3,1,2,2,3,0,0,1]
status=['validé modèle','réfuté prédictif','établi graphe','fortement appuyé','validé modèle','réfuté','non réussi','non testé']
colors=['#3B82F6','#C0392B','#D97706','#D97706','#3B82F6','#C0392B','#C0392B','#98A2B3']
ax.bar(range(len(labels)),values,color=colors)
ax.set_xticks(range(len(labels)),labels,fontsize=8)
ax.set_ylim(0,3.5)
ax.set_yticks([0,1,2,3],['réfuté / nul','preuve de concept','appuyé / interne','validé dans modèle'])
ax.set_title('Statut scientifique consolidé des principales couches')
ax.grid(axis='y',alpha=.2)
for i,(v,s) in enumerate(zip(values,status)):
    ax.text(i,v+0.08,s,ha='center',va='bottom',fontsize=7,rotation=0)
fig.tight_layout()
status_fig=ASSETS/'11_statuts_consolides.png'
fig.savefig(status_fig,dpi=220,bbox_inches='tight')
plt.close(fig)

# ---------- Annex data ----------
# La table de correspondance canonique est maintenue dans le socle.
map_df = pd.read_csv(BASE / '00_socle/genealogie/correspondance_GM_GA.csv', sep=';', encoding='utf-8-sig')
with (ANNEX / 'correspondance_GM_GA.csv').open(
        'w', encoding='utf-8-sig', newline='') as flux:
    map_df.to_csv(flux, sep=';', index=False, lineterminator='\n')

results_rows=[
('Socle','Test interventionnel','11/11 dans le modèle réduit, six cinétiques robustes','Validé dans le modèle réduit'),
('Socle','Effet contre-intuitif','10 cas sur 600, compétition et retard','Exploratoire'),
('Carte','Prédiction de liens','AUC structurelle 0,491, chronologie 0,922','Résultat négatif'),
('Matière','Six dimensions','0,000 bit propre','Codage actuel réfuté comme mesure indépendante'),
('Matière','Échelle de capacités','0,595 bit net, p = 5e-5, rho = 0,74','Établi dans le graphe'),
('Matière','Inventaire accessible','C, H, N recouvrent les coefficients, S en désaccord','Fortement appuyé, tension identifiée'),
('Matière','Mémoire matérielle réelle',
 f"{len(material_memory['transversalite']['familles_soutenantes'])} familles positives sur au moins une relation ; "
 f"{material_memory['transversalite']['familles_au_schema_complet']} chaîne complète sous quatre contrôles, 3 exigées",
 'C-MAT-MEM-05 ne soutient pas la transversalité'),
('Astronomie','N-corps','13 critères sur 15','Validé dans le modèle réduit'),
('Climat','M2 contre M1P','0 critère sur 5, gain RMSE -31,6 %','Réfuté'),
('Exoplanètes','Dépendance au chemin','4 variables sur 4, ablation 4 sur 4','Validation structurelle'),
('Exoplanètes','Persistance','0 variable matérielle après palier long','Non réussi'),
('Matière','Fermeture stricte v0.9.3','46/53 nœuds ; noyau N029-N030-N053-N054 ; réparation candidate 53/53 non canonique','Verrou localisé, source causale insuffisante'),
('Matière','Calibrage structurel v0.9.4','31 nœuds stables, 15 sensibles, 7 dans le verrou canonique ; 40 hyperarêtes à effet d’ablation','Tri structurel et documentaire, causalité générale non démontrée'),
('Climat','Transfert orbital intermédiaire','gain positif dans 3 fenêtres sur 3, 3,12 % en moyenne','Hors échantillon à un pas, non GCM'),
('Mémoire','Bassins et hystérèse','deux bassins pour M2 et M2P ; aucun écart matériel après retour complet','Exploratoire, non confirmatoire'),
('Vivant','Benchmark externe Card 2019','histoire moins bonne dans les 4 groupes ; IC bootstrap défavorable','Externe rétrospectif, non confirmatoire'),
('Vivant','Programme prébiotique','deux trajectoires ARN sur huit cycles, aucune lignée parent-descendant','Critère héréditaire non testé'),
('Généalogie','Clôture du graphe linéaire','39 transitions, 77 relations parent-produit, aucune anomalie formelle','Cohérence structurelle vérifiée, distincte de l’hypergraphe strict')]
with (ANNEX / 'resultats_scientifiques_consolides.csv').open(
        'w', encoding='utf-8-sig', newline='') as flux:
    pd.DataFrame(
        results_rows,
        columns=['domaine','test_ou_objet','resultat','statut'],
    ).to_csv(flux, index=False, lineterminator='\n')

# ---------- DOCX Helpers ----------
DOCX = OUT / 'DOSSIER_SCIENTIFIQUE_ORI-C.docx'
PDF = OUT / 'DOSSIER_SCIENTIFIQUE_ORI-C.pdf'

doc=Document()
sec=doc.sections[0]
sec.page_height=Cm(29.7); sec.page_width=Cm(21.0)
sec.top_margin=Cm(1.7); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)

# theme fonts
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos')
styles['Normal'].font.size=Pt(9.5)
styles['Normal'].paragraph_format.space_after=Pt(5)
styles['Normal'].paragraph_format.line_spacing=1.08
for sname,size,color in [('Title',30,NAVY),('Heading 1',18,NAVY),('Heading 2',13,PURPLE),('Heading 3',10.5,TEAL)]:
    st=styles[sname]
    st.font.name='Aptos Display' if sname!='Normal' else 'Aptos'
    st._element.rPr.rFonts.set(qn('w:eastAsia'),st.font.name)
    st.font.size=Pt(size)
    st.font.color.rgb=RGBColor.from_string(color.replace('#',''))
    st.font.bold=True
    st.paragraph_format.space_before=Pt(10 if sname=='Heading 1' else 7)
    st.paragraph_format.space_after=Pt(5)
    st.paragraph_format.keep_with_next=True
styles['Heading 1'].paragraph_format.page_break_before=True

# utility formatting

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill.replace('#',''))


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)


def set_row_cant_split(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:cantSplit'); trPr.append(el)


def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run('Page '); run.font.size=Pt(8); run.font.color.rgb=RGBColor(100,100,100)
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='PAGE'
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def setup_header_footer(section):
    section.header.is_linked_to_previous=False
    section.footer.is_linked_to_previous=False
    hp=section.header.paragraphs[0]
    hp.clear()
    hp.text='ORI-C  |  Dossier scientifique complet'
    hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(110,110,120)
    fp=section.footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)

setup_header_footer(sec)


def add_para(text='', style=None, bold_start=None, italic=False, align=None, keep=False):
    p=doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        r=p.add_run(bold_start); r.bold=True
        p.add_run(text[len(bold_start):])
    else:
        r=p.add_run(text)
        r.italic=italic
    if align is not None: p.alignment=align
    if keep: p.paragraph_format.keep_with_next=True
    return p


def add_annex_heading(text):
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(8)
    p.paragraph_format.keep_with_next=True
    r=p.add_run(text)
    r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor.from_string(NAVY.replace('#',''))
    return p


def add_callout(title, body, color=PURPLE, fill='#F5F3FF'):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); set_cell_shading(c,fill); set_cell_margins(c,120,160,120,160)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3)
    r=p.add_run(title); r.bold=True; r.font.color.rgb=RGBColor.from_string(color.replace('#','')); r.font.size=Pt(11)
    p2=c.add_paragraph(body); p2.paragraph_format.space_after=Pt(0); p2.paragraph_format.line_spacing=1.06
    return t


def add_figure(path, caption, width_cm=16.8):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path),width=Cm(width_cm))
    cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after=Pt(6)
    for r in cp.runs: r.italic=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(90,90,100)


def add_table(headers, rows, widths=None, font_size=7.5, header_fill=NAVY, repeat=True):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style='Table Grid'
    hdr=table.rows[0]
    for i,h in enumerate(headers):
        cell=hdr.cells[i]; cell.text=str(h); set_cell_shading(cell,header_fill); set_cell_margins(cell,60,60,60,60)
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(font_size)
    if repeat: set_repeat_table_header(hdr)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells; set_row_cant_split(table.rows[-1])
        for i,val in enumerate(row):
            cells[i].text='' if val is None or (isinstance(val,float) and math.isnan(val)) else str(val)
            if ri%2: set_cell_shading(cells[i],'F7F8FA')
            set_cell_margins(cells[i],45,55,45,55); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: r.font.size=Pt(font_size)
        if widths:
            for i,w in enumerate(widths): cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return table


def new_landscape():
    s=doc.add_section(start_type=2) # new page
    s.orientation=WD_ORIENT.LANDSCAPE
    s.page_width=Cm(29.7); s.page_height=Cm(21.0)
    s.top_margin=Cm(1.4); s.bottom_margin=Cm(1.3); s.left_margin=Cm(1.4); s.right_margin=Cm(1.4)
    setup_header_footer(s)
    return s


def new_portrait():
    s=doc.add_section(start_type=2)
    s.orientation=WD_ORIENT.PORTRAIT
    s.page_height=Cm(29.7); s.page_width=Cm(21.0)
    s.top_margin=Cm(1.7); s.bottom_margin=Cm(1.6); s.left_margin=Cm(1.8); s.right_margin=Cm(1.8)
    setup_header_footer(s)
    return s

# ---------- Cover ----------
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55)
r=p.add_run('LE POSSIBLE A UNE HISTOIRE'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(NAVY.replace('#',''))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Dossier scientifique complet ORI-C'); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string(PURPLE.replace('#',''))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Socle théorique, résultats, généalogie, découvertes, réfutations et programme de validation'); r.font.size=Pt(12); r.font.color.rgb=RGBColor(70,75,90)
doc.add_paragraph().paragraph_format.space_before=Pt(18)
add_figure(program_fig,'Architecture générale du programme ORI-C',width_cm=15.2)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20)
r=p.add_run('Didier Daloze  |  Dossier consolidé au 2 août 2026'); r.font.size=Pt(10); r.font.color.rgb=RGBColor(90,90,100)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Dossier scientifique consolidé à partir du socle, des branches, des données, des tests et de la généalogie.'); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=RGBColor(100,100,110)
doc.add_page_break()

# Reader note
add_para('Note de lecture',style='Heading 1')
add_callout('Objet du dossier',
            'Ce document réunit la totalité du périmètre scientifique contenu dans le socle, les trois branches, les données, les tests et la généalogie intégrée. Il ne réduit pas ORI-C à l’arbre. Il distingue le socle théorique, les quatre représentations de la chronologie, les données mesurées, les tests positifs, les résultats négatifs, les programmes non testés et les corrections nécessaires avant publication.',NAVY,'#EEF3FB')
add_para('Le dossier source contient plusieurs objets qui ne sont pas équivalents. La carte TR décrit quarante transitions et quarante-sept relations. L’hypergraphe représente les processus multi-entrées et multi-sorties des branches matière et planétaire. La généalogie GM détaille vingt et une transitions matérielles. L’arbre GA relie trente-huit transitions réparties entre matière, planète et vivant. Les nombres ne doivent jamais être additionnés comme s’ils décrivaient le même niveau de découpage.')
add_para('La règle d’autorité suivie est celle du dossier canonique. Les rapports générés et les fichiers JSON priment sur les textes rédigés. ETAT_DES_PREUVES.md fixe les statuts scientifiques transversaux. La généalogie est intégrée au socle et les corrections du ²⁶Al, du statut du flux neutronique, du calcul d’information et de la grille de certitude sont documentées dans leurs fichiers propres.')

add_para('Sommaire',style='Heading 2')
contents=[
'1. Question scientifique et périmètre',
'2. Architecture scientifique du socle',
'3. Les quatre représentations de l’histoire',
'4. Branche matière et architectures planétaires',
'5. Branche Système solaire, Terre et mémoire historique',
'6. Branche vivant et programme prébiotique',
'7. Arbre généalogique global',
'8. Résultats et découvertes réellement obtenus',
'9. Limites, corrections intégrées et points non démontrés',
'10. Programme de validation prioritaire',
'Annexes. Catalogues, correspondances, critères et sources']
for item in contents:
    p=doc.add_paragraph(style='List Number'); p.add_run(item)

# 1
add_para('1. Question scientifique et périmètre',style='Heading 1')
add_para('ORI-C étudie la manière dont l’histoire physique transforme les possibilités futures d’un système. La proposition ne consiste pas à dire que le passé influence le présent, fait déjà admis dans de nombreux domaines. Elle cherche à décomposer matériellement cette influence en composition, configuration, interactions, environnement, modes de persistance, mémoires, seuils, pertes et chemins encore accessibles.')
add_callout('Formulation centrale',
            'L’histoire physique réorganise la matière. Cette réorganisation modifie son accessibilité, construit des frontières et des réservoirs, crée des gradients, ouvre certaines transformations et en ferme d’autres. Lorsque certaines architectures entretiennent et transmettent leurs propres conditions de persistance, un régime vivant devient possible.',PURPLE,'#F5F3FF')
add_para('Le programme est organisé en un socle commun et trois branches autonomes. Le socle partage un langage et des critères. Il ne partage pas les mécanismes ni les niveaux de preuve. Une réussite numérique dans la branche astronomique ne valide rien dans la branche prébiotique. Une réfutation de la mémoire climatique M2 reste localisée à cette implémentation. Cette séparation est une condition de solidité du dossier.')
add_figure(program_fig,'Figure 1. Socle commun et circulation contrôlée entre les trois branches.',width_cm=15.8)
add_para('Le périmètre couvre la chronologie des architectures matérielles, les filtrages planétaires, la dynamique du Système solaire, une couche de mémoire historique, une application climatique séparée et la branche du vivant. Les dimensions cognitives et les autres développements formels d’ORI-C restent hors de ce dossier.')

# 2
add_para('2. Architecture scientifique du socle',style='Heading 1')
add_para('2.1 Une architecture ne se réduit pas à son inventaire',style='Heading 2')
add_para('Le socle définit une architecture matérielle comme un ensemble de constituants dont la configuration, les interactions et l’environnement produisent une unité ou un régime collectif identifiable. La composition est nécessaire, mais elle ne suffit pas. Deux systèmes contenant les mêmes constituants peuvent présenter des propriétés, des stabilités et des trajectoires différentes lorsque leur organisation change.')
add_callout('État architectural', 'A(t) = [ n(t), G(t), I(t), E(t), Π(t), H(t) ]',PURPLE,'#F5F3FF')
add_table(['Dimension','Contenu','Question opératoire'],[
('n - composition','Inventaire des constituants','Qu’est-ce qui est présent et en quelle quantité ?'),
('G - configuration','Relations spatiales et topologiques','Comment les constituants sont-ils disposés et reliés ?'),
('I - interactions','Liaisons, réactions et flux','Quelles transformations peuvent effectivement se produire ?'),
('E - environnement','Ressources, contraintes et domaine local','Quel milieu rend les transformations possibles ou les empêche ?'),
('Π - persistance','Mode de maintien dans le temps','Par quelle structure ou quel flux l’organisation dure-t-elle ?'),
('H - histoire','Inscriptions héritées et dépendance au chemin','Quelles traces du passé modifient encore la réponse ?')],font_size=8)
add_para('Les propriétés observables sont attribuées à l’architecture entière, Y(t) = Φ[n, G, I, E, Π, H]. Les tests du socle ont toutefois montré que le premier remplissage des six dimensions ne constituait pas six mesures indépendantes. Les codages étaient entièrement expliqués par le régime. Leur information propre était nulle. Cette réfutation porte sur le codage actuel, pas sur l’utilité conceptuelle des dimensions.')

add_para('2.2 Séparer l’état, les mémoires et l’architecture',style='Heading 2')
add_para('Le cadre distingue l’état présent S à une échelle ℓ, les inscriptions héritées m et l’architecture A qui rend la réponse possible. Une perturbation peut seulement déplacer l’état. Elle peut aussi laisser une mémoire ou modifier les composants, les paramètres structurels, les relations et l’opérateur d’évolution lui-même.')
add_callout('Dynamique couplée',
            'Sₜ₊₁^ℓ = F^ℓ[Aₜ^ℓ](Sₜ^ℓ, Uₜ^ℓ, mₜ^ℓ, ξₜ^ℓ)\n'
            'mₜ₊₁^ℓ = 𝒢ₘ^ℓ(mₜ^ℓ, Sₜ^ℓ, Aₜ^ℓ, Uₜ^ℓ, ξₜ^ℓ)\n'
            'Aₜ₊₁^ℓ = Q^ℓ(Aₜ^ℓ, mₜ^ℓ, Sₜ^ℓ, Uₜ^ℓ, ξₜ^ℓ)\n'
            'Pₜ₊₁,ℓ^(s) = P^(s)(Aₜ₊₁^ℓ, mₜ₊₁^ℓ, Cₜ₊₁^ℓ, T, ε)',NAVY,'#EEF3FB')
add_para('Une variation appartient à l’état lorsqu’elle peut être représentée sans modifier l’opérateur d’évolution. Elle devient architecturale lorsqu’il faut modifier les composants, les relations, les paramètres structurels ou l’opérateur. Ce partage dépend du niveau de description et du plancher de bruit, qui doivent être déclarés.')
add_figure(base_layers_fig,'Figure 2. Les cinq couches de l’architecture scientifique ORI-C.',width_cm=15.5)

add_para('2.3 Mémoire distribuée',style='Heading 2')
add_para('L’histoire n’est pas représentée par une mémoire scalaire unique. Elle est distribuée entre plusieurs compartiments, m(t) = [m₁, m₂, …, mₖ], dont les constantes de temps et les mécanismes d’inscription diffèrent. En approximation linéaire, chaque composante peut être décrite par une convolution du forçage passé avec un noyau propre. Le socle précise que ces noyaux et leurs couplages dépendent eux-mêmes de l’état et de l’architecture.')
add_para('Cette distinction est essentielle pour séparer une persistance durable d’un simple retard de relaxation. Le test exoplanétaire l’a démontré négativement. Deux histoires différentes produisaient un écart à court terme, mais celui-ci disparaissait lorsque le palier final commun dépassait les constantes de temps lentes. Le modèle avait une mémoire transitoire, pas une inscription durable.')

add_para('2.4 Persistance et héritage',style='Heading 2')
add_table(['Mode de persistance','Mécanisme','Exemple de régime'],[
('Liée','Liaison ou barrière énergétique','Noyau, molécule, réseau cristallin'),
('Métastable','Maintien hors équilibre sans flux entretenu','Phase persistante derrière une barrière'),
('Dissipative','Entretien par un flux','Structure ouverte alimentée'),
('Homéostatique','Régulation interne','Maintien d’une variable dans une plage'),
('Reconstructive','Réparation et renouvellement','Organisation individuelle capable de se restaurer'),
('Reproductive','Propagation de l’organisation','Transmission entre individus ou générations'),
('Évolutive','Variation héréditaire et sélection','Transformation cumulative d’une lignée')],font_size=8)
add_para('Le seuil du vivant n’est pas présenté comme un degré supérieur sur une échelle unique. Il correspond à un changement de mode de continuité. Une organisation participe activement à la conservation, à la reconstruction, à la reproduction et à la transformation de ses propres conditions de persistance.')
add_table(['Type d’héritage','Ce qui est transmis'],[
('Matériel','Constituants, éléments, réservoirs et produits antérieurs'),
('Configuratif','Topologie, défauts, séquences et organisation spatiale'),
('Dynamique','Cycles, gradients et couplages installés'),
('Contraignant','Voies rendues accessibles, coûteuses ou impossibles'),
('Génératif','Machinerie capable de reproduire et de réinterpréter l’héritage')],font_size=8)

add_para('2.5 Signature des transitions et fermeture des possibilités',style='Heading 2')
add_callout('Signature', 'S = (ΔV, ΔC, ΔΠ, ΔH, ΔR, ΔF)',ORANGE,'#FFF6EA')
add_table(['Terme','Transformation décrite'],[
('ΔV','Apparition de variables collectives nouvelles'),('ΔC','Modification de la connectivité entre états accessibles'),
('ΔΠ','Apparition d’un nouveau mode de persistance'),('ΔH','Importance nouvelle de l’héritage'),
('ΔR','Robustesse, récupération et coût du retour'),('ΔF','Possibilités fermées, contractées ou rendues plus coûteuses')],font_size=8)
add_para('Une transition n’est donc pas seulement une production. Elle peut ouvrir de nouveaux chemins tout en séquestrant des constituants, en supprimant des retours ou en rendant certains états trop coûteux dans l’horizon considéré. La généalogie encode dix-neuf fermetures dans l’arbre global. Elles devront être divisées en fermetures physiques, historiques, pratiques et probabilistes.')

add_para('2.6 D-H-L et domaines de possibles',style='Heading 2')
add_table(['Diagnostic','Mesure','Ce qu’il ne faut pas confondre'],[
('D - durée','Temps de relaxation, de résidence ou de reconstruction','Une mémoire longue n’est pas nécessairement irréversible'),
('H - hystérésis','Écart entre seuil de basculement et seuil de retour','Une asymétrie de retour n’implique pas une perte de composant'),
('L - perte','Disparition d’un composant, d’une relation ou d’un chemin','Une perte peut être lente et sans basculement brutal')],font_size=8)
add_para('À une échelle ℓ, le socle distingue quatre filtres emboîtés : P^adm, compatible avec les lois et contraintes ; P^att, atteignable depuis l’état courant sous le générateur déclaré ; P^kin, accessible avant l’horizon T avec les vitesses, ressources et barrières disponibles ; P^pers, dont la réalisation laisserait une trace au-dessus du seuil de persistance déclaré. Les anciennes notations restent compatibles : Pth correspond à P^adm et Pacc(T,C,ε) à P^kin lorsque l’état initial et le générateur sont inclus dans C.')
add_para('La mesure de persistance Π_pers,t^ℓ = P_pers^ℓ[h_t ; O, W] doit déclarer l’observable O, la fenêtre W et un seuil Π*. La réalisation effective reste séparée des quatre ensembles. La chaîne physique du système doit également être distinguée de la chaîne épistémique D + M → grandeurs inférées.')
add_figure(chain_fig,'Figure 3. Chaîne d’organisation ORI-C. Elle structure l’analyse, mais aucune branche ne la mesure encore de bout en bout.',width_cm=16.2)

add_para('2.7 Relations causales et niveaux de preuve',style='Heading 2')
add_table(['Code','Sens','Portée'],[
('MATR','Fournit les constituants','Causalité matérielle historique'),('ENBL','Rend possible','Condition causale historique'),
('ENVR','Modifie le milieu','Transformation environnementale'),('STAB','Stabilise','Maintien d’une architecture'),
('CATL','Catalyse','Accélère une transformation'),('CNST','Contraint','Réduit le domaine accessible'),
('CONT','Contribue','Non suffisant'),('FEED','Rétroaction','Seul lien cyclique admis dans la carte'),
('DEPG, INCO, DESC','Dépendance, trace, ascendance','Non causal historiquement'),('CLOS','Ferme ou contracte','Défini, pas entièrement instancié'),
('INTG','Intégration durable','Défini pour les événements comme l’endosymbiose')],font_size=7.5)
add_para('Les niveaux Établi, Fortement inféré, Plausible et Hypothétique doivent être séparés du mode de preuve, observation, reconstruction, simulation ou expérimentation. La généalogie distingue désormais quatre axes indépendants, preuve du mécanisme, preuve en milieu naturel, preuve de la transition historique et certitude du rôle causal. Leur évaluation actuelle reste préliminaire et demande une revue bibliographique indépendante.')

# 3 representations
add_para('3. Les quatre représentations de l’histoire',style='Heading 1')
add_figure(repr_fig,'Figure 4. Les quatre représentations successives ne décrivent pas le même objet.',width_cm=16.5)
add_table(['Représentation','Unité','Résultat principal','Limite'],[
('Carte des transitions','40 transitions et 47 relations','Chronologie documentée et liens typés','Topologie dominée par l’ordre chronologique'),
('Hypergraphe mécanistique','53 nœuds et 53 hyperarêtes','Processus multi-entrées, recyclages, séparation matière-condition','Ne mesure pas encore les flux et probabilités'),
('Généalogie détaillée','22 transitions de matière','Parents, produits, preuves, ouvertures et fermetures explicites','Ne couvre pas les branches planétaire et vivant'),
('Arbre généalogique global','39 transitions en trois branches','Clôture formelle et raccordements interbranches','Correspondance explicitée, références primaires encore incomplètes')],font_size=8)

add_para('3.1 Carte des transitions',style='Heading 2')
add_para('La carte historique répartit quarante transitions en huit régimes de cinq transitions. Elle contient quarante-sept relations typées. Seize relations sont classées Établies, douze Fortement inférées, dix-sept Plausibles et deux Hypothétiques. L’analyse de graphe montre une densité de 0,030, cinq communautés, une modularité de 0,659 et un plus long chemin de vingt-quatre transitions.')
add_para('Le test décisif est négatif. Un prédicteur structurel de liens masqués obtient une AUC de 0,491, au niveau du hasard. La proximité chronologique atteint 0,922. La carte enregistre les relations, mais sa topologie n’apporte pas encore une information prédictive indépendante de l’ordre des transitions.')

add_para('3.2 Hypergraphe mécanistique',style='Heading 2')
add_para(f"L’hypergraphe contient {hyper_valid['nodes']} nœuds et {hyper_valid['hyperedges']} hyperarêtes. Il compte {hyper_valid['multi_input_edges']} processus à plusieurs entrées, {hyper_valid['multi_output_edges']} à plusieurs sorties, {hyper_valid['recycling_or_transport_edges_with_shared_endpoint']} processus de transport ou recyclage à point partagé et {hyper_valid['explicit_scenario_or_hypothesis_edges']} scénarios explicitement séparés des faits. Tous les nœuds sont joignables depuis la racine baryonique déclarée.")
add_para('La clôture a révélé une rupture réelle dans la première construction. La chaîne poussière se recyclait sans alimentation depuis les condensats, les glaces, les molécules organiques et les grains présolaires. Une hyperarête a dû être ajoutée pour rétablir la continuité matérielle. Ce contrôle est un résultat de méthode. Il détecte des filiations manquantes que la chronologie narrative ne rendait pas visibles.')
add_para('Le reclassement des quarante-sept relations montre que treize seulement correspondent à une filiation matérielle stricte. Vingt-trois sont des conditions d’ouverture, cinq des transformations environnementales, deux des contraintes d’inventaire, trois des dépendances non généalogiques et une une transmission de trace historique.')

add_para('3.3 Généalogies détaillée et globale',style='Heading 2')
add_para(f"La généalogie détaillée contient {close_gm['transitions']} transitions, {close_gm['produits_distincts']} produits distincts et {close_gm['relations_parent_produit']} relations parent-produit. Elle déclare {close_gm['possibilites_fermees_declarees']} possibilités fermées. L’arbre global contient {close_ga['transitions']} transitions, {close_ga['produits_distincts']} produits distincts et {close_ga['relations_parent_produit']} relations parent-produit. Les trois branches comptent {close_ga['transitions_par_branche']['1 matière']}, {close_ga['transitions_par_branche']['2 planétaire']} et {close_ga['transitions_par_branche']['3 vivant']} transitions.")
add_para('Le CSV global contient quatre-vingt-une mentions de parents. Soixante-dix-sept correspondent à des produits déjà formés dans l’arbre. Quatre sont traitées comme entrées externes : quarks, gluons et deux occurrences des électrons. Les neutrons sont désormais distingués du flux neutronique, qui est codé comme condition dynamique.')

# 4 matter/planet
add_para('4. Branche matière et architectures planétaires',style='Heading 1')
add_para('4.1 Les huit régimes de la base historique',style='Heading 2')
reg_rows=[]
for (num,name),grp in tr.groupby(['regime_num','regime_nom']):
    reg_rows.append((num,name,' → '.join(grp['transition'].tolist())))
add_table(['Régime','Nom','Transitions historiques'],reg_rows,font_size=7.2)
add_para('Cette base constitue un inventaire historique. Les six dimensions y sont renseignées qualitativement, mais neuf champs essentiels restent vides dans toutes les lignes, notamment les preuves directes, les modèles concurrents, les seuils, les vitesses de variation, les pertes, les mécanismes de persistance et les contre-exemples. Leur remplissage exige de la littérature primaire et des évaluateurs indépendants.')

add_para('4.2 Généalogie matérielle détaillée',style='Heading 2')
add_figure(branch_figs['1 matière'],'Figure 5. Branche matérielle de l’arbre global. Le statut de chaque transition est indiqué dans son encadré.',width_cm=16.4)
add_para('La branche matérielle suit la formation des nucléons, noyaux, atomes, molécules, étoiles, éléments lourds, grains, glaces, organiques et agrégats. Son apport principal n’est pas la découverte de ces étapes, qui appartiennent aux disciplines établies. Il réside dans leur représentation comme une chaîne de produits réutilisables, accompagnée des conditions, mécanismes, preuves, possibilités ouvertes et fermées.')

add_para('4.3 Échelle des dix capacités',style='Heading 2')
add_para('L’hypergraphe propose dix capacités de travail, depuis les constituants baryoniques jusqu’aux interfaces et gradients entretenus. Le test de monotonie est réfuté. Douze baisses inattendues subsistent dans la première règle et une violation subsiste même après révision exploratoire. Une étoile organisée à un niveau élevé produit des éléments simples. Un système hydrothermal produit des espèces mobiles de niveau inférieur. L’échelle ordonne des objets, pas les processus qui les produisent.')
add_para(f"La non-redondance résiste toutefois au test par permutation. L’échelle porte {hierarchy['test_B_information']['gain_net_du_bruit_bits']} bit net, avec p = {hierarchy['test_B_information']['p_par_permutation']:.0e} et une corrélation de Spearman de {hierarchy['test_B_information']['spearman_niveau_profondeur']:.2f} avec la profondeur. Ce résultat est interne au graphe. Il valide un codage informatif, pas une mesure naturelle universelle.")

add_para('4.4 Inventaire accessible',style='Heading 2')
add_callout('Définition opérationnelle', 'Inventaire accessible = quantité × fraction mobilisable × probabilité de transfert × min(1, durée disponible / horizon)',TEAL,'#EAF8F4')
inv_rows=[]
for el in ['N','C','H','S']:
    d=inventory['par_element'][el]
    lo,hi=d['etendue_de_la_part_accessible']
    inv_rows.append((el,d['reservoir_de_reference'],f'{lo:.3f} à {hi:.3f}',f"× {d['facteur_d_indetermination']:.1f}",f"{inventory['classement_par_sequestration_maximale'][['N','C','H','S'].index(el)][1]:.2f}"))
add_table(['Élément','Réservoir de référence','Part accessible selon scénarios','Indétermination','Rapport maximal noyau/référence'],inv_rows,font_size=8)
add_para('Les trente et un enregistrements couvrent le carbone, l’hydrogène, l’azote et le soufre. Les budgets noyau plus silicate reconstituent les totaux publiés avec un écart maximal de 2,88 %. Les coefficients de partage métal-silicate indépendants recouvrent les répartitions observées pour le carbone, l’hydrogène et l’azote. Le soufre reste en désaccord. Cette tension reproduit un désaccord publié entre une estimation classique riche en soufre et des coefficients expérimentaux indiquant un noyau plus pauvre.')
add_para('Ce résultat mesure une répartition entre réservoirs. Il ne mesure pas encore les fractions mobilisables, les flux de transfert ni Pacc à un horizon donné. Il fournit néanmoins la première mise en données réelle de la distinction entre inventaire total et inventaire disponible dans le dossier.')

add_para('4.5 Campagne de mémoire matérielle réelle',style='Heading 2')
mm = material_memory['transversalite']
add_para(f"La campagne WP-MAT-MEM-2026 trouve {len(mm['familles_soutenantes'])} familles positives sur au moins une relation, mais {mm['familles_au_schema_complet']} famille porte la chaîne complète histoire → trace → réponse sous les quatre contrôles conjoints, alors que trois étaient exigées. Le verdict C-MAT-MEM-05 est donc {mm['verdict']}. Les résultats relationnels locaux ne doivent pas être transformés en validation transversale.")
add_para('Un test même état apparent, même stimulus et histoires différentes est ajouté comme analyse exploratoire. Son plan ayant été choisi après inspection de la table, il ne rend aucun verdict confirmatoire et sert uniquement à préparer une réplication indépendante préenregistrée.')

add_para('4.6 Filtrages historiques planétaires',style='Heading 2')
add_table(['Filtrage','Conséquence architecturale','Statut disciplinaire'],[
('Lieu de formation','Sélection des matériaux et des signatures isotopiques','Établi'),
('Moment de l’accrétion','Quantité de ²⁶Al restante et destin thermique','Établi'),
('Différenciation métal-silicate','Création du noyau, du manteau et séquestration des éléments','Établi'),
('Dégazage et pertes','Réorganisation des volatils et de l’atmosphère','Établi'),
('Apports tardifs','Modification d’une architecture déjà différenciée','Établi'),
('Fusion et impacts','Effacement partiel ou recomposition de la mémoire initiale','Fortement appuyé')],font_size=8)
add_para('Cette chaîne est le lieu où la dépendance au chemin est matériellement enregistrée dans la base. Elle est portée par la cosmochimie, la radiochronologie et la pétrologie expérimentale. ORI-C l’organise, mais ne la découvre pas. La valeur ajoutée propre du cadre sur cette chaîne reste à tester.')
add_figure(branch_figs['2 planétaire'],'Figure 6. Branche planétaire de l’arbre global.',width_cm=16.4)

# 5 solar
add_para('5. Branche Système solaire, Terre et mémoire historique',style='Heading 1')
add_para('5.1 Couche astronomique',style='Heading 2')
add_para('La couche astronomique repose sur vingt-cinq calculs N-corps, une trajectoire principale de vingt millions d’années, des conditions initiales JPL Horizons DE441 et des comparaisons à Horizons et La2010. Treize critères préenregistrés sur quinze sont réussis. Le modèle réduit est validé numériquement et astronomiquement dans son périmètre.')
astro_rows=[
('Tous les corps liés','vrai','vrai','réussi'),('Conservation de l’énergie','1,325 × 10⁻¹¹','≤ 10⁻⁸','réussi'),('Moment angulaire newtonien','5,276 × 10⁻¹⁰','≤ 10⁻¹⁰','échec'),
('Excentricité initiale contre La2010','2,192 × 10⁻¹⁰','≤ 10⁻⁸','réussi'),('Corrélation Horizons à 6 ka','≈ 1','≥ 0,99','réussi'),('RMSE Horizons à 6 ka','4,826 × 10⁻⁷','≤ 2 × 10⁻⁴','réussi'),
('Corrélation La2010 à 100 ka','0,999971','≥ 0,95','réussi'),('Corrélation La2010 à 500 ka','0,998760','≥ 0,80','réussi'),('Corrélation La2010 à 1 Ma','0,997270','≥ 0,60','réussi'),
('Convergence du pas sur 2 Ma','8,427 × 10⁻⁷','≤ 10⁻⁴','réussi'),('WHFast contre IAS15 à 20 ka','3,131 × 10⁻⁷','≤ 10⁻⁶','réussi'),('Aller-retour à 100 ka','2,763 × 10⁻⁵','≤ 10⁻⁵','échec'),
('Pic spectral de 405 ka','0,007861','≤ 0,05','réussi'),('Pic spectral de 2,4 Ma','0,166625','≤ 0,20','réussi'),('Contrefactuels au-dessus du plancher','6,27 × 10⁶','≥ 3','réussi')]
add_table(['Critère','Observé','Seuil','Verdict'],astro_rows,font_size=7.5)
add_para('Les interventions sur Jupiter et Saturne produisent des écarts de plusieurs millions de fois supérieurs à la dispersion d’états initiaux presque identiques. Cela établit un effet causal interne au modèle de l’architecture gravitationnelle. Le modèle ne résout toutefois ni la Lune, ni la rotation terrestre, ni le J₂ solaire, ni les marées, ni l’obliquité dynamique. Il ne produit aucune prédiction climatique ou géologique hors échantillon.')

add_para('5.2 Couche mémoire historique et MPT',style='Heading 2')
add_para('La couche mémoire testait une implémentation précise de l’inscription climatique sur la transition du Pléistocène moyen. Le protocole initial a été corrigé après identification de cinq défauts, notamment l’absence de témoin de complexité égale, l’usage du BIC sur des résidus fortement autocorrélés, des paramètres aux bornes et un palier exoplanétaire plus court que les constantes de temps lentes.')
add_table(['Comparaison','Critères réussis','Gain de RMSE','Intervalle 95 %','ΔBIC effectif'],[
('M2 contre M1, témoin moins complexe','1 / 5','+3,57 %','[2,72 %, 4,57 %]','+5,46'),
('M2 contre M1P, complexité égale','0 / 5','−31,56 %','[−38,93 %, −25,08 %]','+9,37')],font_size=8)
add_para('Le modèle M2 produit un rapport spectral 100/41 ka de 0,0047, alors que LR04 atteint 2,604 dans la fenêtre de prédiction. Le couplage carbone dégrade la RMSE hors échantillon après ablation contrôlée. Les paramètres de mémoire restent aux bornes ou non identifiés. La déclinaison M2 est réfutée et doit rester abandonnée afin d’éviter un ajustement post hoc.')

add_para('5.3 Test exoplanétaire contrôlé',style='Heading 2')
add_para('Deux histoires spin-orbitales différentes aboutissent au même forçage final. Le modèle avec mémoire produit une dépendance au chemin significative sur quatre variables et l’ablation ramène les écarts au niveau nul sur quatre variables. Le test structurel du code est donc réussi.')
add_para('Aucune variable ne franchit toutefois le seuil d’amplitude physique défini avant le calcul. Après un palier final de trois cents millions d’années, l’écart s’annule. Le temps caractéristique de décroissance est d’environ sept millions d’années. Le résultat correspond à un retard de relaxation, pas à une inscription durable.')

add_para('5.4 Application climatique séparée',style='Heading 2')
add_para('L’article climatique est hors de la chaîne de preuve de la branche. Il ne reçoit aucun verdict des tests N-corps ou MPT. Son apport au socle réside dans des distinctions devenues transversales : mémoire distribuée, D-H-L, hiérarchie des possibles, séparation S/m/A, seuil de persistance, critère d’altération architecturale et séparation des chaînes physique et épistémique. Les contenus propres au climat restent des éléments de littérature et une étude de cas, pas une validation d’ORI-C.')

# 6 living
add_para('6. Branche vivant et programme prébiotique',style='Heading 1')
add_para('6.1 Statut de la branche',style='Heading 2')
add_table(['Objet','Statut','Portée'],[
('Cellule eucaryote comme architecture','Preuve de concept','Application descriptive des dimensions'),
('Endosymbiose mitochondriale','Preuve de concept','Faits biologiques fortement appuyés, représentation ORI-C non testée'),
('Résistance aux antibiotiques','Non testé dans la branche canonique','Protocole expérimental proposé'),
('Programme prébiotique','Critère héréditaire non testé','Trajectoires ARN réelles, aucune lignée de compartiments'),
('Universalité et pouvoir prédictif','Non testé','Aucune validation transversale')],font_size=8)
add_figure(branch_figs['3 vivant'],'Figure 7. Branche du vivant de l’arbre global.',width_cm=16.4)

add_para('Les données Papastavrou ajoutent deux trajectoires expérimentales de populations d’ARN catalytique suivies pendant huit cycles. Elles mesurent une dynamique de composition, sans relier des compartiments parents à leurs descendants. Le critère de continuité héréditaire reste non testé.')

add_para('6.2 Le verrou prébiotique',style='Heading 2')
add_para('Le programme ne situe pas le verrou dans une brique isolée comme l’ARN, la membrane ou une réaction métabolique. Il le situe dans le couplage entre compartimentation, copie par matrice, variation héritable, apport énergétique et persistance sur plusieurs cycles. Une molécule produite n’est pas une hérédité. Une fonction transitoire sans information transmise reste un état. Une information transmise sans effet fonctionnel reste une trace.')
add_table(['Axe','Donnée décisive'],[
('Formation des briques','Rendements, sous-produits, stabilité et vitesse'),('Polymérisation','Longueurs, séquences, dégradation et répétabilité'),
('Copie par matrice','Fidélité, erreurs, blocages et dépendance à la séquence'),('Compartimentation','Encapsulation, perméabilité, croissance, division et durée'),
('Couplage copie-compartiment','Rétention et partage des produits à la division'),('Apport énergétique','Gradients, énergie libre, consommation et rendement'),
('Variation héritable','Transmission des variantes sur plusieurs cycles'),('Sélection prébiotique','Différences de croissance, copie, résistance et division'),
('Dépendance au chemin','Ordres différents de chauffage, séchage, gel ou irradiation'),('Intégration en protocellule','Maintien simultané de membrane, copie, énergie et transmission')],font_size=7.5)
add_para('Le critère minimal exige six conditions simultanées. Une matrice produit des copies avec variation. Les copies restent liées à un compartiment. Les compartiments croissent et se divisent. Des variantes sont transmises. Certaines variantes changent la persistance ou la reproduction. Cette différence se maintient plusieurs cycles sans réinitialisation complète par l’expérimentateur.')
add_para('La donnée centrale manque encore. Elle serait la proportion de compartiments descendants conservant à la fois une information moléculaire héritée et une différence fonctionnelle mesurable après plusieurs cycles. Sans table de lignées, le programme observe une production chimique, pas une hérédité.')

add_para('6.3 Boucle de retour du vivant',style='Heading 2')
add_para('Le vivant n’est pas seulement le terme aval de la branche planétaire. Il transforme l’atmosphère, les sols, les minéraux, les cycles et les sédiments. Il produit de nouvelles inscriptions géologiques et modifie le domaine de possibilités de l’architecture terrestre. Cette boucle est conceptuellement essentielle, mais elle n’est pas encore intégrée quantitativement dans l’arbre généalogique.')

# 7 la généalogie tree
add_para('7. Arbre généalogique global',style='Heading 1')
add_para('7.1 Résultats formels',style='Heading 2')
add_table(['Indicateur','Valeur'],[
('Transitions',str(close_ga['transitions'])),
('Branche matière',str(close_ga['transitions_par_branche']['1 matière'])),
('Branche planétaire',str(close_ga['transitions_par_branche']['2 planétaire'])),
('Branche vivant',str(close_ga['transitions_par_branche']['3 vivant'])),
('Produits distincts',str(close_ga['produits_distincts'])),
('Relations parent-produit internes',str(close_ga['relations_parent_produit'])),
('Transferts interbranches comptés','9'),
('Couples de transitions interbranches distincts','8'),
('Transitions établies',str(close_ga['certitudes_synthetiques']['établi'])),
('Fortement inférées',str(close_ga['certitudes_synthetiques']['fortement inféré'])),
('Plausibles',str(close_ga['certitudes_synthetiques']['plausible'])),
('Fermetures déclarées',str(close_ga['possibilites_fermees_declarees'])),
('Anomalies formelles',str(len(close_ga['anomalies'])))],font_size=8.2)
add_para('La clôture signifie que chaque parent matériel est produit auparavant ou déclaré comme entrée externe, que l’ordre est respecté et que les vocabulaires autorisés sont employés. Elle démontre la cohérence du graphe avec ses règles de construction. Elle ne démontre ni l’exactitude des mécanismes, ni la qualité des preuves citées, ni la réalisation historique de chaque transition.')

add_figure(tree_fig,'Figure 8. Vue d’ensemble de l’arbre généalogique global et liste des huit couples de transitions interbranches.',width_cm=16.4)

add_para('7.2 Transferts interbranches',style='Heading 2')
add_figure(inter_fig,'Figure 9. Les huit couples de transitions distincts qui assurent la continuité matérielle entre les branches.',width_cm=16.4)
add_para('Le graphe contient neuf transferts matériels interbranches correspondant à huit couples de transitions distincts. Le couple GA-029 vers GA-030 transporte deux produits, les minéraux secondaires et les gradients redox. Les autres couples transportent chacun un produit déclaré.')

add_para('7.3 Ce que l’arbre permet désormais',style='Heading 2')
add_para('L’arbre localise l’origine déclarée de chaque produit, les conditions permissives, les voies concurrentes, les ouvertures, les fermetures et les zones de faible certitude. Il transforme une chronologie en objet interrogeable. On peut retirer un nœud, suivre ses dépendances aval, repérer les goulets d’étranglement, comparer des histoires et préparer des tests d’ablation ou de contrefactualité.')
add_callout('Résultat démontré par la généalogie',
            'Une représentation généalogique formellement close et calculable peut relier les transformations matérielles, planétaires et biologiques dans une structure commune. Ce résultat est une validation de construction. Il ne constitue pas encore la généalogie démontrée du réel ni une preuve de supériorité prédictive.',TEAL,'#EAF8F4')

# 8 findings
add_para('8. Résultats et découvertes réellement obtenus',style='Heading 1')
add_figure(status_fig,'Figure 10. Statut consolidé des couches principales. Les hauteurs ne forment pas une échelle universelle de preuve.',width_cm=16.5)
add_para('8.1 Découvertes d’architecture scientifique',style='Heading 2')
architecture_findings=[
('La transition comme unité causale complète','Parents matériels, conditions, mécanisme, produit, propriété, accessibilité, ouvertures, fermetures et preuve sont réunis dans une même unité.'),
('Séparation matière-condition','Ce dont un produit est fait est distingué de ce qui rend sa formation possible.'),
('Accessibilité plutôt qu’inventaire total','L’histoire agit en redistribuant les constituants entre réservoirs, surfaces, interfaces et voies de transfert.'),
('Fermeture du futur','Une transformation peut créer des capacités tout en supprimant des retours, en séquestrant des éléments ou en rendant des états trop coûteux.'),
('Mémoire distribuée','Plusieurs compartiments conservent le passé avec des constantes de temps et des couplages différents.'),
('Passage au vivant','Le changement de régime se situe dans l’entretien et la transmission des conditions de persistance, pas dans une simple augmentation de complexité.'),
('Goulets interbranches','Quelques transferts matériels concentrent la continuité entre matière, planète et vivant.'),
('Incertitude localisée','Le graphe montre exactement quelles articulations sont établies, inférées ou plausibles.')]
add_table(['Résultat structurel','Contenu'],architecture_findings,font_size=8)

add_para('8.2 Résultats propres aux tests du dossier',style='Heading 2')
add_table(['Domaine','Objet','Résultat','Statut'],results_rows,font_size=7.2)
add_para('Les résultats négatifs font partie des découvertes du programme. La carte des transitions ne prédit pas au-delà de la chronologie. Les six dimensions ne portent aucune information propre dans leur premier codage. La monotonie de l’échelle des capacités est réfutée. M2 ne bat pas un témoin de complexité égale. La dépendance exoplanétaire disparaît sur palier long. Ces échecs définissent les implémentations à abandonner et les critères à durcir.')

add_para('8.3 Ce qui appartient aux disciplines et ce qui appartient à ORI-C',style='Heading 2')
add_table(['Élément','Origine scientifique','Apport ORI-C actuel'],[
('Nucléosynthèse, chimie interstellaire, grains','Physique, astrophysique, cosmochimie','Réorganisation généalogique et typage des relations'),
('Accrétion, différenciation, volatils','Sciences planétaires et pétrologie','Lien entre histoire des réservoirs et inventaire accessible'),
('Dépendance au chemin planétaire','Cosmochimie et radiochronologie','Intégration dans une architecture commune'),
('Membranes, polymères, endosymbiose','Chimie prébiotique et biologie','Décomposition du verrou d’intégration et protocole de lignées'),
('Test N-corps','Calcul produit dans le dossier','Validation du modèle réduit'),
('Test interventionnel','Calcul produit dans le dossier','Théorème local et limites structurelles'),
('Clôture généalogique','Calcul produit dans le dossier','Validation formelle de la généalogie encodée')],font_size=8)
add_para('ORI-C n’a pas encore découvert une loi physique ou biologique universelle. Sa découverte actuelle est une architecture de représentation et de test. Elle devient scientifiquement forte lorsqu’elle produit un résultat discriminant que les descriptions classiques de complexité égale ne produisent pas.')

add_para('8.4 Campagne ciblée v0.9.3',style='Heading 2')
add_table(['Priorité exécutée','Résultat','Portée'],[
('Verrou matière','Le noyau cyclique est localisé sur N029, N030, N053 et N054. Un recodage candidat ferme 53 nœuds sur 53.','La source disponible ne démontre pas la direction causale exacte. La réparation reste hors hypergraphe canonique.'),
('Transfert orbital-climat','Le signal N-corps améliore la RMSE dans trois fenêtres temporelles sur trois, de 3,12 % en moyenne.','Prédiction à un pas utilisant l’état climatique observé. LR04 est accordée orbitalement. Aucun GCM.'),
('Mémoire durable','M2 et son témoin apparié présentent deux bassins et une boucle d’hystérèse à 30 degrés.','Aucun état matériellement différent ne subsiste après le retour complet. Résultat exploratoire.'),
('Réplication vivant','Dans Card 2019, le modèle historique est moins bon dans les quatre groupes de test.','Jeu externe aux données Windels, mais analyse rétrospective et non confirmatoire.'),
('Prébiotique','Deux trajectoires réelles de populations d’ARN catalytique sont intégrées sur huit cycles.','Aucune filiation parent-descendant de compartiments. La continuité héréditaire reste non testable.')],font_size=7.4)
add_para('Cette campagne transforme les priorités en objets exécutables, mais aucun résultat ne reçoit un statut plus large que son protocole. La fermeture candidate n’est pas injectée dans le graphe canonique, le transfert climatique n’est pas un GCM, l’hystérèse reste interne au modèle et les deux benchmarks biologiques ne constituent pas une confirmation prospective.')

add_para('8.5 Calibrage structurel et documentaire v0.9.4',style='Heading 2')
add_table(['Objet calibré','Résultat','Interprétation et limite'],[
('Graphe canonique','Les 53 nœuds et 53 hyperarêtes de la v0.9.3 sont gelés et contrôlés par empreinte.','Le calibrage n’ajoute ni ne retire de relation canonique.'),
('Ablations','40 hyperarêtes sur 53 produisent une perte mesurable dans la projection ou la fermeture stricte.','Une importance structurelle dans le graphe ne démontre pas une nécessité causale dans la nature.'),
('Stress documentaire','Sur 4 000 tirages déterministes, 31 nœuds sont stables, 15 sensibles, 0 fragile et 7 déjà bloqués par le verrou canonique.','Les fréquences mesurent la dépendance au codage documentaire actuel, pas une probabilité de vérité.'),
('Tri des relations','35 priorités d’ablation, 4 relations du cycle de verrou, 10 cycles d’entretien, 3 relations redondantes et 1 priorité documentaire à fort effet aval.','H011, l’instabilité de streaming, devient la première cible documentaire hors module des interfaces.'),
('Seuils documentaires','Le profil complet atteint 53 nœuds en projection et 46 en fermeture stricte. Le profil équilibré atteint 44 et 31.','La contraction localise la dépendance aux relations moins documentées. Elle ne les réfute pas.'),
('Benchmark externe MESA','Deux trajectoires stellaires, 12 transitions et 14 nœuds atteignent une fermeture stricte de 14 sur 14.','Transfert du schéma de représentation, sans validation observationnelle de MESA ni loi universelle ORI-C.')],font_size=7.2)
add_para('Le calibrage apporte un tri plus fin entre soutien documentaire, rôle structurel, redondance, dépendance cyclique et effet en aval. Il conserve explicitement comme non mesurées la nécessité empirique, la suffisance, la temporalité quantitative, la réversibilité physique et l’effet d’une intervention directe. Il sert donc à choisir les prochains tests, pas à transformer le graphe en preuve causale générale.')

# 9 limits corrections
add_para('9. Limites, corrections intégrées et points non démontrés',style='Heading 1')
add_para('9.1 La clôture est formelle',style='Heading 2')
add_para('Le vérificateur confirme la provenance interne des parents, la présence des champs, l’ordre des transitions, les catégories de mécanisme et les valeurs de certitude. Il ne lit pas la littérature et ne teste pas la causalité historique. La formulation correcte est que le graphe est cohérent avec ses règles. La formulation incorrecte serait qu’il démontre la généalogie réelle de la matière au vivant.')

add_para('9.2 Le calcul d’information de la généalogie GM',style='Heading 2')
add_para('Un ancien calcul d’information annonçait un gain conditionnel maximal de 0,654 bit. Cette valeur a été retirée des résultats probants et conservée uniquement dans les archives non probantes. Les champs textuels identifiaient presque chaque transition et permettaient une mémorisation des lignes. Le protocole actuel exige des catégories indépendantes de la cible, une validation hors échantillon et des permutations.')
add_para('Une mesure recevable exige des catégories définies indépendamment de la cible, une séparation apprentissage-validation, des permutations, des groupes de transitions laissés hors échantillon et une comparaison à la position chronologique seule.')

add_para('9.3 Correction scientifique du ²⁶Al',style='Heading 2')
add_para('La filiation du ²⁶Al est maintenant séparée de la branche des actinides et du processus r. Le ²⁶Al est produit dans les étoiles massives par plusieurs sites de combustion, notamment le cœur convectif en combustion de l’hydrogène, la coquille convective du carbone et la combustion explosive néon-carbone. Il est ensuite éjecté, injecté dans le milieu présolaire ou le disque, incorporé aux solides et agit comme source de chauffage précoce des planétésimaux.')
add_figure(al26_fig,'Figure 11. Filiation corrigée proposée pour le ²⁶Al. Les actinides restent dans une branche distincte de nucléosynthèse lourde.',width_cm=16.3)
add_para('Références externes de correction : C. Iliadis et al., The Astrophysical Journal Supplement Series 193, 16, 2011, DOI 10.1088/0067-0049/193/1/16. M. Limongi et A. Chieffi, The Astrophysical Journal 647, 483-500, 2006, DOI 10.1086/505164. I. D. Hutcheon et R. Hutchison, Nature 337, 238-241, 1989, DOI 10.1038/337238a0.')

add_para('9.4 Neutrons et flux neutronique',style='Heading 2')
add_para('Les généalogies distinguent désormais les noyaux graines et les neutrons comme parents matériels, tandis que le flux neutronique lent ou rapide est codé comme condition permissive du mécanisme.')

add_para('9.5 Raccordement GM-GA',style='Heading 2')
add_para('La correspondance entre la généalogie détaillée et l’arbre global est désormais explicitée dans une table canonique. Elle indique les correspondances directes, les étapes ajoutées dans l’arbre global et le passage de la différenciation détaillée vers plusieurs produits planétaires.')

add_para('9.6 Certitude multidimensionnelle',style='Heading 2')
add_table(['Colonne proposée','Question'],[
('preuve_du_mecanisme','Le mécanisme est-il produit, mesuré ou reproduit ?'),
('preuve_en_milieu_naturel','Le mécanisme existe-t-il dans un environnement naturel pertinent ?'),
('preuve_de_la_transition_historique','Cette transition précise a-t-elle eu lieu dans la trajectoire considérée ?'),
('certitude_du_role_causal','Le mécanisme a-t-il joué le rôle causal attribué dans cette transition ?')],font_size=8)
add_para('Cette séparation empêche d’attribuer à une voie historique un statut élevé simplement parce que le mécanisme existe en laboratoire. Une synthèse de type Strecker peut être établie expérimentalement sans démontrer qu’elle a produit les premiers monomères biologiques sur la Terre primitive.')

add_para('9.7 Reproductibilité technique',style='Heading 2')
add_para('La version 0.9.4-research aligne la source, le verrou de dépendances, les manifestes, le calibrage matière et les workflows de publication. Elle ajoute une CI séparée sous Python 3.12 et 3.13, un contrôle strict de Git LFS, une comparaison inter-environnements avec tolérances numériques explicites et une archive canonique déterministe. Les piles numériques natives sont exécutées dans des processus séparés pour éviter les blocages d’orchestration. La publication reste conditionnée au passage effectif des workflows après dépôt et création du tag.')

# 10 program
add_para('10. Programme de validation prioritaire',style='Heading 1')
add_para('Le plan directeur contient huit cent quatre items répartis en cinquante-huit groupes de travail. Le présent dossier réduit ce programme à une séquence de priorités qui conditionne directement la valeur scientifique du cadre.')
add_table(['Priorité','Action','Critère de réussite'],[
('1','Revoir bibliographiquement le ²⁶Al, le flux neutronique et la correspondance GM-GA','Corrections intégrées puis validées par une revue indépendante'),
('2','Calibrer et documenter les relations prioritaires, notamment H011 et le cycle H030-H031-H052-H053','Séparation explicite entre soutien documentaire, rôle structurel et causalité empirique'),
('3','Sourcer chaque transition à la littérature primaire','Référence identifiable pour mécanisme, milieu, histoire et rôle causal'),
('4','Réviser les quatre axes de certitude transition par transition','Aucun transfert implicite du laboratoire vers l’histoire réelle'),
('5','Appliquer le protocole d’information généalogique','Validation hors échantillon et permutation contre chronologie'),
('6','Mesurer l’inventaire accessible réel','Flux, spéciation, probabilités et horizons, pas seulement répartition'),
('7','Tester deux histoires à inventaire total voisin','Différence prédite par la distribution des réservoirs et vérifiée sur données'),
('8','Préenregistrer une prédiction ORI-C minimale','Témoin de complexité égale, seuil, ablation et arrêt définis avant le résultat'),
('9','Durcir le programme prébiotique','Lignées, témoin non copiable apparié et persistance sur plusieurs cycles'),
('10','Répliquer à l’extérieur','Au moins un résultat positif et un résultat négatif reproduits indépendamment'),
('11','Réparer la chaîne logicielle','Source, wheel, verrou, conteneur et manifeste issus de la même révision')],font_size=7.5)
add_para('Le test décisif pour ORI-C reste le même. Il faut montrer que l’histoire, l’accessibilité ou l’architecture apportent une prédiction hors échantillon qu’un modèle classique de complexité égale, connaissant la composition et l’état présent, ne produit pas. Tant que ce résultat manque, le cadre doit être présenté comme un programme falsifiable en consolidation.')

add_para('Conclusion générale',style='Heading 1')
add_para('La totalité du dossier montre une progression en trois mouvements. Le socle a construit un langage commun, des modèles, des données et des procédures de réfutation. Le programme a obtenu un théorème local dans le chémostat, une validation N-corps dans un modèle réduit, une première mesure de l’inventaire accessible et plusieurs résultats négatifs qui éliminent des implémentations fragiles. La campagne v0.9.3 localise le verrou de fermeture stricte de la matière, construit un transfert orbital-climat intermédiaire, teste les bassins et l’hystérèse, exécute un benchmark antibiotique externe et intègre des trajectoires d’ARN réelles. Le calibrage v0.9.4 gèle ensuite cette architecture, mesure les effets d’ablation, distingue 31 nœuds stables de 15 nœuds sensibles et transfère le schéma à deux trajectoires stellaires MESA. Le graphe généalogique linéaire reste formellement clos selon ses règles, tandis que l’hypergraphe mécanistique strict demeure ouvert à 46 nœuds sur 53.')
add_para('L’architecture scientifique propre à ORI-C se situe dans l’articulation entre matière transmise, conditions permissives, organisation, mémoire, persistance, accessibilité et fermeture des possibilités. Elle ne constitue pas encore une loi générale démontrée. Elle constitue une méthode structurée permettant de dire exactement ce qui est connu, ce qui est seulement organisé, ce qui est réfuté et ce qui doit encore être testé.')
add_callout('Verdict scientifique consolidé',
            'ORI-C possède désormais un socle conceptuel cohérent, plusieurs objets calculables, des résultats disciplinaires et numériques réels, une généalogie formellement close et une capacité croissante à produire ses propres réfutations. Il ne possède pas encore une prédiction transversale positive, reproduite et supérieure à un témoin apparié. La prochaine avancée dépend moins de l’élargissement du vocabulaire que d’un test discriminant réussi.',NAVY,'#EEF3FB')

# Annexes
new_landscape()
add_annex_heading('Annexe A. Catalogue compact des 39 transitions de l’arbre global')
rows=[]
for _,r in ga.iterrows():
    rows.append((r['id'],r['branche'],r['parents_materiels'],r['conditions_permissives'],r['produit'],r['mecanisme_categorie'].replace('_',' '),r['possibilites_ouvertes'],r['possibilites_fermees'],r['degre_de_certitude']))
headers_ga=['ID','Branche','Parents matériels','Conditions','Produit','Mécanisme','Ouvertures','Fermetures','Certitude']
ga_chunks=[rows[:13],rows[13:26],rows[26:]]
for idx,chunk in enumerate(ga_chunks):
    if idx:
        doc.add_page_break()
        add_annex_heading(f'Annexe A. Catalogue de l’arbre global - suite {idx+1}/3')
    add_table(headers_ga,chunk,font_size=6.1,header_fill=NAVY)
new_portrait()

add_para('Annexe B. Correspondance entre GM et GA',style='Heading 1')
add_table(list(map_df.columns),map_df.itertuples(index=False,name=None),font_size=7.5)

new_landscape()
add_annex_heading('Annexe C. Catalogue des 40 transitions historiques')
tr_rows=[]
for _,r in tr.iterrows():
    tr_rows.append((r['id'],r['regime_num'],r['regime_nom'],r['transition'],r['date'],r['etats_devenus_accessibles'],r['etats_fermes'],r['niveau_de_preuve']))
headers_tr=['ID','Régime','Nom du régime','Transition','Fenêtre','États accessibles','États fermés','Niveau']
tr_chunks=[tr_rows[:14],tr_rows[14:27],tr_rows[27:]]
for idx,chunk in enumerate(tr_chunks):
    if idx:
        doc.add_page_break()
        add_annex_heading(f'Annexe C. Catalogue des transitions historiques - suite {idx+1}/3')
    add_table(headers_tr,chunk,font_size=6.0)
new_portrait()

add_para('Annexe D. Résultats consolidés',style='Heading 1')
add_table(['Domaine','Objet','Résultat','Statut'],results_rows,font_size=7.5)

add_para('Annexe E. Sources internes faisant autorité',style='Heading 1')
source_rows=[
('Architecture et vocabulaire','00_socle/CODEBOOK.md, ARCHITECTURE.md'),
('Statuts scientifiques','ETAT_DES_PREUVES.md'),
('Compteurs logiciels','ETAT_DES_TESTS.md'),
('Carte relationnelle','00_socle/carte_relationnelle/ANALYSE_GRAPHE.md'),
('Test interventionnel','00_socle/test_interventionnel/PORTEE_WP_S2.md et rapports générés'),
('Branche matière','01_branche_matiere/README.md et base_transitions/'),
('Hypergraphe et inventaire','01_branche_matiere/hypergraphe_transformations/*.json et *.csv'),
('Filtrages planétaires','02_branche_systeme_solaire/FILTRAGES_HISTORIQUES.md'),
('Astronomie','couche_astronomique/STATUT_SCIENTIFIQUE.md et VALIDATION_FINALE.md'),
('Mémoire historique','couche_memoire_historique/REPORT.md, RAPPORT_CORRIGE.md et STRESS_REPORT.md'),
('Vivant et prébiotique','03_branche_vivant/README.md et PROGRAMME_PREBIOTIQUE.md'),
('Campagne ciblée v0.9.3','plan_directeur/campagne_priorites_v093/ et rapports des cinq paquets'),
('Calibrage matière v0.9.4','01_branche_matiere/hypergraphe_transformations/calibrage_v094/ et protocole gelé v0.9.3'),
('Benchmark externe','03_branche_vivant/benchmark_externe_card2019/ et DOI 10.5061/dryad.g41hg96'),
('Trajectoires ARN','03_branche_vivant/programme_prebiotique/donnees_reelles/ et DOI 10.5061/dryad.rxwdbrvgs'),
('Arbre généalogique','00_socle/genealogie/arbre_genealogique.csv et cloture_arbre.json'),
('Généalogie détaillée de la matière','01_branche_matiere/genealogie/genealogie_matiere.csv et cloture_genealogie.json')]
add_table(['Domaine','Fichiers'],source_rows,font_size=8)

add_para('Annexe F. Livrables du présent dossier',style='Heading 1')
add_para('Le dossier de livraison contient le document Word, sa version PDF, les figures utilisées, les tables de correspondance GM-GA, la synthèse des résultats et des copies des principaux fichiers machine lisibles. Les données, codes, rapports et annexes du dossier complet restent disponibles dans leur arborescence d’origine.')

# metadata
props=doc.core_properties
props.title='Le possible a une histoire - Dossier scientifique complet ORI-C'
props.subject='Socle, branches, résultats, généalogie, découvertes et programme de validation'
props.author='Didier Daloze'
props.keywords='ORI-C, histoire, architecture, généalogie, matière, planète, vivant, mémoire, persistance'

# save
for p in doc.paragraphs:
    if p.style.name=='Normal':
        p.paragraph_format.widow_control=True

doc.save(DOCX)

# README dossier
readme=f"""# Dossier scientifique ORI-C

Ce dossier regroupe :

- `{DOCX.name}` : version modifiable
- `{PDF.name}` : version de lecture
- `annexes/` : tables et fichiers machine lisibles
- `assets/` : figures intégrées au document

Le dossier rassemble le socle, les branches, les résultats, les données et la généalogie intégrée. Les résultats probants, exploratoires, négatifs et non testés restent explicitement distingués.
"""
(OUT/'README.md').write_text(readme,encoding='utf-8')
print(DOCX)
