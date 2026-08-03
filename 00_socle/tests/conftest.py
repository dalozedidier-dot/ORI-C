"""Fixtures partagées de la suite de vérification ORI-C."""
from __future__ import annotations

import csv
import hashlib
from pathlib import Path
import shutil
import subprocess
import sys

import pytest

RACINE = Path(__file__).resolve().parent.parent
DOSSIER = RACINE.parent
CARTE = RACINE / "carte_relationnelle"
INTERVENTION = RACINE / "test_interventionnel"
ARTICLE = DOSSIER / "01_branche_matiere" / "article"
SCRIPTS_INTERVENTION = INTERVENTION / "scripts"

sys.path.insert(0, str(SCRIPTS_INTERVENTION))


def lire_csv(chemin: Path) -> list[dict[str, str]]:
    """Les CSV du dossier sont en UTF-8 avec BOM et séparateur point-virgule."""
    with chemin.open(encoding="utf-8-sig", newline="") as flux:
        return list(csv.DictReader(flux, delimiter=";"))


def sha256(chemin: Path) -> str:
    return hashlib.sha256(chemin.read_bytes()).hexdigest()


def sha256_normalise(chemin: Path) -> str:
    """Empreinte insensible aux fins de ligne, pour comparer entre systèmes."""
    return hashlib.sha256(chemin.read_bytes().replace(b"\r\n", b"\n")).hexdigest()


def graphviz_disponible() -> bool:
    return shutil.which("dot") is not None


besoin_graphviz = pytest.mark.skipif(
    not graphviz_disponible(),
    reason="L'exécutable Graphviz `dot` est absent du PATH ; la régénération de la carte est impossible.",
)


@pytest.fixture(scope="session")
def racine() -> Path:
    return RACINE


@pytest.fixture(scope="session")
def noeuds() -> list[dict[str, str]]:
    return lire_csv(CARTE / "data/noeuds_poc.csv")


@pytest.fixture(scope="session")
def liens() -> list[dict[str, str]]:
    return lire_csv(CARTE / "data/relations_oric_47_provisoires.csv")


@pytest.fixture(scope="session")
def exoplanetes() -> list[dict[str, str]]:
    return lire_csv(CARTE / "data/cas_exoplanetes.csv")


@pytest.fixture(scope="session")
def article_docx():
    from docx import Document

    return Document(ARTICLE / "Chronologie_des_architectures_de_la_matiere_ORI-C.docx")


@pytest.fixture(scope="session")
def article_texte(article_docx) -> str:
    """Paragraphes et cellules de tableau, pour ne rien manquer."""
    morceaux = [paragraphe.text for paragraphe in article_docx.paragraphs]
    for tableau in article_docx.tables:
        for ligne in tableau.rows:
            morceaux.extend(cellule.text for cellule in ligne.cells)
    return "\n".join(morceaux)


@pytest.fixture(scope="session")
def metriques_intervention(racine: Path) -> dict:
    """Relance le test interventionnel publié, puis relit ses métriques.

    La relance garantit que l'on teste le script et non un résidu sur disque.
    """
    import json

    subprocess.run(
        [sys.executable, str(SCRIPTS_INTERVENTION / "test_interventionnel_ori_c.py")],
        cwd=racine, check=True, capture_output=True,
    )
    chemin = INTERVENTION / "resultats/metriques_test_interventionnel_ori_c.json"
    return json.loads(chemin.read_text(encoding="utf-8"))


@pytest.fixture(scope="session")
def manifeste(racine: Path) -> list[tuple[str, str]]:
    entrees = []
    for ligne in (DOSSIER / "MANIFEST.sha256").read_text(encoding="utf-8").splitlines():
        if not ligne.strip():
            continue
        empreinte, _, relatif = ligne.partition("  ")
        entrees.append((empreinte.strip(), relatif.strip()))
    return entrees
